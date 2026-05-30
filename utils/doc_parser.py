"""
Document parser: extract raw text from PDF and DOCX files.
"""

from pathlib import Path
from loguru import logger


def extract_text(file_path: Path) -> str:
    """
    Extract plain text from a PDF or DOCX file.
    Returns empty string on failure (caller decides how to handle).
    """
    suffix = file_path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _extract_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return _extract_docx(file_path)
        else:
            logger.warning("Unsupported file type: {}", suffix)
            return ""
    except Exception as e:
        logger.error("Failed to extract text from {}: {}", file_path.name, e)
        return ""


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    text = "\n".join(pages).strip()

    if not text:
        # Scanned PDF — OCR would be needed; flag it
        logger.warning("PDF '{}' appears to be a scanned image with no extractable text. "
                       "Consider adding OCR (e.g. pytesseract) for scanned documents.", path.name)
    return text


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)
